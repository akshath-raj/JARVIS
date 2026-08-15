"""Author a well-formatted Pages document using its REAL built-in paragraph styles.

Pages has no scripting API for its paragraph styles (the Title / Subtitle / Heading
picker in the Format sidebar), and that control isn't reachable by UI scripting on
current builds. The robust, deterministic way to get genuine Pages styles is to
build a Word ``.docx`` with the standard named styles and open it in Pages: Pages
imports Word's built-in styles and maps them onto its own — ``Title`` → Pages Title
(with its underline rule), ``Heading 1`` → Pages Heading, ``Normal`` → Body, etc. —
so the document opens with real, editable Pages styles in the picker (verified).

This authors from a simple list of blocks:

    [{"type": "title",    "text": "Deepfakes"},
     {"type": "subtitle", "text": "An Overview"},
     {"type": "heading",  "text": "What Are Deepfakes?", "level": 1},
     {"type": "body",     "text": "They are synthetic media…"},
     {"type": "image",    "path": "/…/pic.png", "caption": "A deepfake"},
     {"type": "table",    "rows": [["Signal", "Tell"], ["Blinking", "Irregular"]]}]

Image blocks take a LOCAL ``path`` (the caller resolves any web images through the
safe fetcher first — this module never touches the network). The built file is then
opened in Pages via LaunchServices.
"""
from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Inches

# Map our block heading levels to Word style names Pages understands.
_HEADING_STYLE = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
_MAX_IMAGE_WIDTH_IN = 6.0   # usable page width on the default template


class AuthorError(RuntimeError):
    pass


class PagesAuthor:
    """Build a styled .docx from blocks and open it in Pages."""

    def create_document(self, blocks: list[dict], out_path: str) -> str:
        """Write `blocks` to a .docx at `out_path` using real named styles. Returns
        the path actually written (extension forced to .docx)."""
        doc = Document()
        for block in blocks:
            self._add_block(doc, block)
        p = Path(out_path).expanduser()
        if p.suffix.lower() != ".docx":
            p = p.with_suffix(".docx")
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
        return str(p)

    def _add_block(self, doc: Document, block: dict) -> None:
        kind = (block.get("type") or "body").strip().lower()
        text = block.get("text", "") or ""
        if kind == "title":
            doc.add_paragraph(text, style="Title")
        elif kind == "subtitle":
            doc.add_paragraph(text, style="Subtitle")
        elif kind == "heading":
            level = int(block.get("level", 1) or 1)
            doc.add_paragraph(text, style=_HEADING_STYLE.get(level, "Heading 1"))
        elif kind == "body":
            # split on blank lines so each paragraph is its own Body paragraph
            for para in (text.split("\n\n") if text else [""]):
                doc.add_paragraph(para.strip(), style="Normal")
        elif kind == "image":
            self._add_image(doc, block)
        elif kind == "table":
            self._add_table(doc, block.get("rows") or [])
        else:  # unknown → treat as body text so nothing is silently lost
            doc.add_paragraph(text, style="Normal")

    def _add_image(self, doc: Document, block: dict) -> None:
        path = block.get("path", "")
        if not path or not Path(path).expanduser().is_file():
            raise AuthorError(f"image file not found: {path!r}")
        doc.add_picture(str(Path(path).expanduser()), width=Inches(_MAX_IMAGE_WIDTH_IN * 0.8))
        caption = block.get("caption", "")
        if caption:
            doc.add_paragraph(caption, style="Caption")

    def _add_table(self, doc: Document, rows: list[list]) -> None:
        rows = [r for r in rows if r]
        if not rows:
            return
        ncols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=ncols)
        table.style = "Table Grid"  # a real bordered style Pages imports cleanly
        for i, row in enumerate(rows):
            for j in range(ncols):
                table.rows[i].cells[j].text = str(row[j]) if j < len(row) else ""

    def open_in_pages(self, path: str) -> None:
        """Open the built document in Pages (LaunchServices launches + focuses it)."""
        try:
            subprocess.run(
                ["open", "-a", "Pages", str(Path(path).expanduser())],
                capture_output=True, text=True, timeout=20, check=True,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
            raise AuthorError("couldn't open the document in Pages") from e
