"""Tests for authoring a Pages document via a .docx with REAL named styles.

We build the .docx and read it back with python-docx to assert each block landed on
the correct Word paragraph style (which Pages maps to its own Title/Heading/Body).
No GUI or network is involved.
"""
from __future__ import annotations

import pytest
from docx import Document

from jarvis.tools.pages_author import AuthorError, PagesAuthor


def _styles(path):
    return [(p.text, p.style.name) for p in Document(path).paragraphs if p.text]


def test_blocks_map_to_real_named_styles(tmp_path):
    blocks = [
        {"type": "title", "text": "Deepfakes"},
        {"type": "subtitle", "text": "An Overview"},
        {"type": "heading", "text": "What Are Deepfakes?", "level": 1},
        {"type": "heading", "text": "Sub point", "level": 2},
        {"type": "body", "text": "First para.\n\nSecond para."},
    ]
    out = PagesAuthor().create_document(blocks, str(tmp_path / "d"))
    assert out.endswith(".docx")  # extension forced
    got = _styles(out)
    assert ("Deepfakes", "Title") in got
    assert ("An Overview", "Subtitle") in got
    assert ("What Are Deepfakes?", "Heading 1") in got
    assert ("Sub point", "Heading 2") in got
    # a body block with a blank line becomes two Body (Normal) paragraphs
    assert ("First para.", "Normal") in got
    assert ("Second para.", "Normal") in got


def test_table_block_creates_grid_table(tmp_path):
    blocks = [{"type": "table", "rows": [["A", "B"], ["1", "2"], ["3", "4"]]}]
    out = PagesAuthor().create_document(blocks, str(tmp_path / "t.docx"))
    doc = Document(out)
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert (len(t.rows), len(t.columns)) == (3, 2)
    assert t.cell(0, 0).text == "A" and t.cell(2, 1).text == "4"
    assert t.style.name == "Table Grid"


def _valid_png_1x1() -> bytes:
    import struct
    import zlib

    def chunk(typ, data):
        body = typ + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)   # 1x1, 8-bit truecolour
    idat = zlib.compress(b"\x00\xff\x00\x00")             # filter byte + one red pixel
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def test_image_block_embeds_file(tmp_path):
    png = tmp_path / "pic.png"
    png.write_bytes(_valid_png_1x1())
    out = PagesAuthor().create_document(
        [{"type": "image", "path": str(png), "caption": "A picture"}],
        str(tmp_path / "img.docx"),
    )
    doc = Document(out)
    assert len(doc.inline_shapes) == 1                       # image embedded
    assert any(p.text == "A picture" for p in doc.paragraphs)  # caption present


def test_missing_image_raises(tmp_path):
    with pytest.raises(AuthorError):
        PagesAuthor().create_document(
            [{"type": "image", "path": "/no/such/file.png"}], str(tmp_path / "x.docx")
        )


def test_unknown_block_kept_as_body(tmp_path):
    out = PagesAuthor().create_document(
        [{"type": "mystery", "text": "kept as body"}], str(tmp_path / "u.docx")
    )
    assert ("kept as body", "Normal") in _styles(out)


def test_open_in_pages_uses_launchservices(tmp_path, monkeypatch):
    import subprocess as sp
    calls = {}
    monkeypatch.setattr(sp, "run",
                        lambda cmd, **k: calls.setdefault("cmd", cmd) or
                        sp.CompletedProcess(cmd, 0, "", ""))
    PagesAuthor().open_in_pages(str(tmp_path / "d.docx"))
    assert calls["cmd"][:3] == ["open", "-a", "Pages"]
